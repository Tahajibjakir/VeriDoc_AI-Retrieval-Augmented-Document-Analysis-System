from typing import Any, List
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from app.api import deps
from app.schemas.project import Project, ProjectCreate, ProjectUpdate
from app.models.models import Project as ProjectModel, User

router = APIRouter()

@router.get("/", response_model=List[Project])
def read_projects(
    db: Session = Depends(deps.get_db),
    current_user: User = Depends(deps.get_current_user),
    skip: int = 0,
    limit: int = 100,
) -> Any:
    """
    Retrieve projects.
    """
    projects = db.query(ProjectModel).filter(ProjectModel.owner_id == current_user.id).offset(skip).limit(limit).all()
    # TODO: Add document_count and questionnaire_count logic
    return projects

@router.post("/", response_model=Project)
def create_project(
    *,
    db: Session = Depends(deps.get_db),
    project_in: ProjectCreate,
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Create new project.
    """
    project = ProjectModel(
        name=project_in.name,
        client_name=project_in.client_name,
        due_date=project_in.due_date,
        scope_type=project_in.scope_type,
        owner_id=current_user.id,
    )
    db.add(project)
    db.commit()
    db.refresh(project)
    return project

@router.get("/{project_id}", response_model=Project)
def read_project(
    *,
    db: Session = Depends(deps.get_db),
    project_id: int,
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Get project by ID.
    """
    project = db.query(ProjectModel).filter(
        ProjectModel.id == project_id,
        ProjectModel.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project

@router.delete("/{project_id}", response_model=Project)
def delete_project(
    *,
    db: Session = Depends(deps.get_db),
    project_id: int,
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Delete a project.
    """
    project = db.query(ProjectModel).filter(
        ProjectModel.id == project_id,
        ProjectModel.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    import os
    from app.core.chroma import get_collection
    collection = get_collection()
    
    # 1. Cleanup physical files and Chroma embeddings for this project's documents
    for doc in project.documents:
        # ChromaDB
        try:
            collection.delete(where={"document_id": doc.id})
        except Exception as e:
            print(f"Error deleting chroma embeddings for doc {doc.id}: {e}")
            
        # Physical File
        if doc.file_path and os.path.exists(doc.file_path):
            try:
                os.remove(doc.file_path)
            except Exception as e:
                print(f"Error deleting physical file {doc.file_path}: {e}")

    # 2. Database cleanup (cascades handle the associated Models)
    db.delete(project)
    db.commit()
    
    return project

from fastapi import BackgroundTasks
from fastapi.responses import StreamingResponse
import pandas as pd
import io
from app.services.jobs import create_job, JobType, run_generation_job
from app.models.models import Question, Answer, AnswerVersion

@router.post("/{project_id}/generate")
def generate_answers(
    *,
    db: Session = Depends(deps.get_db),
    project_id: int,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Trigger background generation of answers for a project's questionnaire.
    """
    project = db.query(ProjectModel).filter(
        ProjectModel.id == project_id,
        ProjectModel.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    job = create_job(db, JobType.GENERATE, project_id)
    background_tasks.add_task(run_generation_job, job.id, project_id)
    
    return {"status": "success", "job_id": job.id, "message": "Generation started"}

@router.get("/{project_id}/export")
def export_report(
    *,
    db: Session = Depends(deps.get_db),
    project_id: int,
    format: str = "csv",
    current_user: User = Depends(deps.get_current_user),
) -> Any:
    """
    Export all questions and current answers as a CSV, Excel, Word, or PDF file.
    """
    project = db.query(ProjectModel).filter(
        ProjectModel.id == project_id,
        ProjectModel.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
        
    questions = db.query(Question).filter(Question.project_id == project_id).order_by(Question.order_index).all()
    
    from app.models.models import Citation, DocumentChunk, Document as DocumentModel
    
    data = []
    for q in questions:
        answer_text = ""
        yes_no = ""
        references = []
        answer = db.query(Answer).filter(Answer.question_id == q.id).first()
        if answer and answer.current_version_id:
            version = db.query(AnswerVersion).filter(AnswerVersion.id == answer.current_version_id).first()
            if version:
                answer_text = version.answer_text
                yes_no = version.answerability or ""
                
                # Fetch citations for Reference column
                citations = db.query(Citation).filter(Citation.answer_version_id == version.id).all()
                for cit in citations:
                    # Logic to get filename
                    chunk = db.query(DocumentChunk).filter(DocumentChunk.id == cit.document_chunk_id).first()
                    if chunk:
                        doc = db.query(DocumentModel).filter(DocumentModel.id == chunk.document_id).first()
                        if doc:
                            references.append(f"[Source: {doc.filename}, Page: {cit.page_number}]")
        
        # Combine explanation with citations
        ref_combined = answer_text
        if references:
            ref_combined += "\n\nReferences:\n" + "\n".join(list(set(references)))
                
        data.append({
            "Question Number": q.question_number,
            "Question": q.question_text,
            "Yes/No": yes_no,
            "Reference": ref_combined
        })
        
    format = format.lower()
    cols = ["Question Number", "Question", "Yes/No", "Reference"]
    
    if format == "csv":
        df = pd.DataFrame(data, columns=cols)
        stream = io.StringIO()
        df.to_csv(stream, index=False)
        response = StreamingResponse(iter([stream.getvalue()]), media_type="text/csv")
        response.headers["Content-Disposition"] = f"attachment; filename=Project_{project_id}_Report.csv"
        return response
        
    elif format == "excel":
        df = pd.DataFrame(data, columns=cols)
        stream = io.BytesIO()
        with pd.ExcelWriter(stream, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        stream.seek(0)
        response = StreamingResponse(iter([stream.getvalue()]), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        response.headers["Content-Disposition"] = f"attachment; filename=Project_{project_id}_Report.xlsx"
        return response
        
    elif format == "word":
        from docx import Document
        doc = Document()
        doc.add_heading(f'Project {project_id} Report', 0)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Question Number'
        hdr_cells[1].text = 'Question'
        hdr_cells[2].text = 'Yes/No'
        hdr_cells[3].text = 'Reference'
        
        for row in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row['Question Number'])
            row_cells[1].text = str(row['Question'])
            row_cells[2].text = str(row['Yes/No'] or '')
            row_cells[3].text = str(row['Reference'] or '')
            
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        response = StreamingResponse(iter([stream.getvalue()]), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        response.headers["Content-Disposition"] = f"attachment; filename=Project_{project_id}_Report.docx"
        return response
        
    elif format == "pdf":
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        
        stream = io.BytesIO()
        doc = SimpleDocTemplate(stream, pagesize=landscape(letter))
        elements = []
        
        styles = getSampleStyleSheet()
        style_normal = styles['Normal']
        
        table_data = [['Question Number', 'Question', 'Yes/No', 'Reference']]
        for row in data:
            q_num = Paragraph(str(row['Question Number']), style_normal)
            q_text = Paragraph(str(row['Question']), style_normal)
            y_n = Paragraph(str(row['Yes/No'] or ''), style_normal)
            ref_text = Paragraph(str(row['Reference'] or '').replace('\n', '<br/>'), style_normal)
            table_data.append([q_num, q_text, y_n, ref_text])
            
        # Adjusted column widths for 4 columns: increased Q# width to 120
        t = Table(table_data, colWidths=[120, 200, 60, 320]) 
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        
        elements.append(t)
        doc.build(elements)
        
        stream.seek(0)
        response = StreamingResponse(iter([stream.getvalue()]), media_type="application/pdf")
        response.headers["Content-Disposition"] = f"attachment; filename=Project_{project_id}_Report.pdf"
        return response
        
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")
